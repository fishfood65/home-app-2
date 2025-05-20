import streamlit as st
import re
from mistralai import Mistral, UserMessage, SystemMessage
from dotenv import load_dotenv
import os
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.text.run import Run
import re
import time
from PIL import Image
import io
import uuid
import json
from bs4 import BeautifulSoup

st.set_page_config(
    page_title="Hello",
    page_icon="👋",
)

st.write("# Welcome to Home Hero Academy! 👋")

st.markdown(
    """
    ### Your Mission
    Accept a series of challenges testing your knowledge about your home to empower you and your deputies to become a heroic guardian for your home and its precious contents
    """
    )
st.markdown(
    """
    ### Start your Training!
    """
    )
# Generate the AI prompt
api_key = os.getenv("MISTRAL_TOKEN")
client = Mistral(api_key=api_key)

if not api_key:
    api_key = st.text_input("Enter your Mistral API key:", type="password")

if api_key:
    st.success("API key successfully loaded.")
else:
   st.error("API key is not set.")

   # Display environment variables in the Streamlit app
#st.title("Environment Variables")

# Display all environment variables
#env_vars = "\n".join([f"{key}: {value}" for key, value in os.environ.items()])
#st.text(env_vars)

# Main entry point of the app

def download_session_state_as_csv():
    """
    Serializes all keys and values in st.session_state into a CSV
    and provides a download button.
    """
    # 1) Build a list of records
    records = []
    for key, val in st.session_state.items():
        # JSON-encode complex objects
        try:
            value_str = json.dumps(val, default=str)
        except Exception:
            value_str = str(val)
        records.append({
            "key": key,
            "value": value_str
        })

    # 2) Create a DataFrame
    df = pd.DataFrame(records)

    # 3) Convert to CSV
    csv = df.to_csv(index=False).encode("utf-8")

    # 4) Offer download
    st.download_button(
        label="📥 Download data as CSV",
        data=csv,
        file_name="home_data.csv",
        mime="text/csv"
    )

PROGRESS_FILE = "user_progress.json"

def load_progress():
    if os.path.exists(PROGRESS_FILE):
        with open(PROGRESS_FILE, "r") as f:
            return json.load(f)
    return {}

def save_progress(progress):
    with open(PROGRESS_FILE, "w") as f:
        json.dump(progress, f)

def main():
    download_session_state_as_csv()
    
    levels = ("Level 1", "Level 2", "Level 3", "Level 4", "Level 5", "Bonus Level")

    # Initialize session state
    if "section" not in st.session_state:
        st.session_state.section = levels[0]
    if "progress" not in st.session_state:
        st.session_state.progress = load_progress()

    # Sidebar navigation
    selected = st.sidebar.radio(
        "Choose a Level:",
        levels,
        index=levels.index(st.session_state.section),
        key="sidebar_level_radio"
    )

    # Progress indicators…
    st.markdown("#### 🧭 Progress")
    completed = sum(
        1 for i in range(1, 6)
        if st.session_state.progress.get(f"level_{i}_completed", False)
    )
    total_levels = 5
    percent_complete = int(completed / total_levels * 100)

    st.progress(percent_complete)

    # (Optional) show numeric fraction
    st.write(f"{completed} of {total_levels} levels completed ({percent_complete}%)")

    # Enforce Level 1 lock
    if selected != "Level 1" and not st.session_state.progress.get("level_1_completed", False):
        st.warning("🚧 Please complete Level 1 before accessing other levels.")
        st.session_state.section = "Level 1"
    else:
        st.session_state.section = selected

    section = st.session_state.section

    # === your existing levels 1–4 ===
    if section == "Level 1":
        st.subheader("🏁 Welcome to Level 1 Home Basics")
        home()
    elif section == "Level 2":
        st.subheader("🚨 Level 2 Emergency Preparedness")
        emergency_kit_utilities()
    elif section == "Level 3":
        st.subheader("📬 Level 3 Mail & Trash Handling")
        mail_trash_handling()
    elif section == "Level 4":
        st.subheader("🏡 Level 4 Home Services")
        security_convenience_ownership()

    # === Level 5: now with st.tabs ===
    elif section == "Level 5":
        st.subheader("💼 Level 5 Critical Documents")
        tabs = st.tabs([
            "📝 Select Documents",
            "📋 Review Selections",
            "🗂 Document Details",
            "📦 Generate Kit"
        ])

        with tabs[0]:
            st.markdown("### Step 1: Pick Critical Documents")
            emergency_kit_critical_documents()

        with tabs[1]:
            st.markdown("### Step 2: Review Your Picks")
            review_selected_documents()

        with tabs[2]:
            st.markdown("### Step 3: Fill in Document Details")
            collect_document_details()
        
        with tabs[3]:
            generate_kit_tab()

    # === Bonus Level ===
    elif section == "Bonus Level":
        st.subheader("🎁 Bonus Level Content")
        bonus_level()

    # Reset button
    if st.sidebar.button("🔄 Reset Progress", key="btn_reset"):
        st.session_state.progress = {}
        save_progress({})
        st.experimental_rerun()

#### Reusable Functions to Generate and Format Runbooks #####
def format_output_for_docx(output: str) -> str:
    """Formats markdown-like output to docx-friendly text with HTML-like formatting."""
    if not output:
        return ""
    
    # Convert markdown-like headings, bold, and italic
    formatted_text = re.sub(r"^## (.*)", r"<h2>\1</h2>", output, flags=re.MULTILINE)  # Convert ## to <h2>
    formatted_text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", formatted_text)  # Convert **bold** to <b>
    formatted_text = re.sub(r"\*(.*?)\*", r"<i>\1</i>", formatted_text)  # Convert *italic* to <i>
    return formatted_text

    return formatted_text

def save_docx_from_formatted_text(formatted_text: str, doc_filename: str, doc_heading: str):
    """Saves the formatted text into a DOCX file with proper formatting."""
    doc = Document()
    doc.add_heading(doc_heading, 0)

    # Split the formatted text by paragraphs and process each paragraph
    paragraphs = formatted_text.split('\n\n')
    for para in paragraphs:
        # Create a paragraph in the DOCX file
        doc_paragraph = doc.add_paragraph()

        # Handle bold and italic text
        runs = re.split(r'(<b>.*?</b>|<i>.*?</i>)', para)  # Split by bold or italic tags
        for run in runs:
            if run.startswith('<b>'):
                run_text = run[3:-4]  # Remove <b> and </b>
                doc_paragraph.add_run(run_text).bold = True
            elif run.startswith('<i>'):
                run_text = run[3:-4]  # Remove <i> and </i>
                doc_paragraph.add_run(run_text).italic = True
            else:
                doc_paragraph.add_run(run)

    doc.save(doc_filename)


def generate_runbook_from_prompt(
    prompt: str,
    api_key: str,
    button_text: str,
    doc_heading: str,
    doc_filename: str
):
    """
    Reusable Streamlit function to handle LLM completion and export a DOCX or HTML file.
    """
    unique_key = f"{button_text.lower().replace(' ', '_')}_button"
    clicked = st.button(button_text, key=unique_key)

    if clicked:
        st.write("✅ Button was clicked")

        if st.session_state.get("user_confirmation") and prompt:
            try:
                st.write("⏳ Sending to Mistral...")

                client = Mistral(api_key=api_key)
                completion = client.chat.complete(
                    model="mistral-small-latest",
                    messages=[SystemMessage(content=prompt)],
                    max_tokens=2000,
                    temperature=0.5,
                )

                output = completion.choices[0].message.content
                st.success(f"{doc_heading} generated successfully!")
                st.write(output)

                # Format the output into HTML-like format
                formatted_output = format_output_for_docx(output)

                # Saving to DOCX
                if doc_filename.endswith(".docx"):
                    save_docx_from_formatted_text(formatted_output, doc_filename, doc_heading)

                    with open(doc_filename, "rb") as f:
                        st.download_button(
                            label="📄 Download DOCX",
                            data=f,
                            file_name=doc_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                # Saving to HTML (Optional for HTML output)
                elif doc_filename.endswith(".html"):
                    html_content = f"<html><head><title>{doc_heading}</title></head><body>{formatted_output}</body></html>"

                    with open(doc_filename, "w") as f:
                        f.write(html_content)

                    with open(doc_filename, "rb") as f:
                        st.download_button(
                            label="📄 Download HTML",
                            data=f,
                            file_name=doc_filename,
                            mime="text/html"
                        )

            except Exception as e:
                st.error(f"❌ Failed to generate runbook: {str(e)}")

        else:
            st.warning("⚠️ Prompt not confirmed or missing.")

def generate_runbook_from_prompt_split(
    prompt_emergency: str,
    prompt_mail_trash: str,
    api_key: str,
    button_text: str,
    doc_heading: str,
    doc_filename: str
):
    """
    Calls Mistral API with two prompts: emergency/utilities and mail/trash.
    Concatenates results and creates a formatted DOCX file.
    """
    unique_key = f"{button_text.lower().replace(' ', '_')}_button"
    clicked = st.button(button_text, key=unique_key)

    # 🔍 Debug Info
    with st.expander("🧪 Debug Info (Prompt + State)"):
        st.write("🔘 **Button Clicked:**", "✅ Yes" if clicked else "❌ No")
        st.write("🙋 **User Confirmed Prompt:**", st.session_state.get("user_confirmation", False))
        st.write("🔑 **API Key Loaded:**", "✅ Yes" if api_key else "❌ No")

    # Check each prompt presence
        st.write("📄 **Emergency Prompt Exists:**", "✅ Yes" if prompt_emergency else "❌ No")
        if prompt_emergency:
            st.code(prompt_emergency[:500] + "..." if len(prompt_emergency) > 500 else prompt_emergency, language="markdown")

        st.write("📬 **Mail & Trash Prompt Exists:**", "✅ Yes" if prompt_mail_trash else "❌ No")
        if prompt_mail_trash:
            st.code(prompt_mail_trash[:500] + "..." if len(prompt_mail_trash) > 500 else prompt_mail_trash, language="markdown")

    # Optionally display selected session state keys
        st.write("📋 **Selected Session State Keys:**")
        st.json({key: st.session_state.get(key) for key in [
            "user_confirmation",
            "electricity_provider",
            "natural_gas_provider",
            "water_provider",
            "internet_provider",
            "emergency_kit_status",
            "emergency_kit_location"
        ]})

    if clicked:
        st.write("✅ Button was clicked")

        if not st.session_state.get("user_confirmation", False):
            st.warning("⚠️ Please confirm the AI prompt before generating the runbook.")
            return

        try:
            client = Mistral(api_key=api_key)

            st.info("📡 Querying Mistral for Emergency & Utilities Section...")
            emergency_response = client.chat.complete(
                model="mistral-small-latest",
                messages=[UserMessage(content=prompt_emergency)],
                max_tokens=1500,
                temperature=0.5,
            ).choices[0].message.content

            st.info("📬 Querying Mistral for Mail & Trash Section...")
            mail_trash_response = client.chat.complete(
                model="mistral-small-latest",
                messages=[UserMessage(content=prompt_mail_trash)],
                max_tokens=1000,
                temperature=0.5,
            ).choices[0].message.content

            # Combine sections and format for DOCX
            full_output = f"{emergency_response}\n\n{mail_trash_response}"
            formatted_output = format_output_for_docx(full_output)

            st.success(f"{doc_heading} generated successfully!")
            st.write(full_output)

            # Write to DOCX
            doc = Document()
            doc.add_heading(doc_heading, 0)
            doc.add_paragraph(formatted_output)
            doc.save(doc_filename)

            # Download button
            with open(doc_filename, "rb") as f:
                st.download_button(
                    label="📄 Download DOCX",
                    data=f,
                    file_name=doc_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"❌ Error generating runbook: {str(e)}")

def generate_runbook_from_multiple_prompts(
    prompts: list,
    api_key: str,
    button_text: str,
    doc_heading: str,
    doc_filename: str
):
    """
    Calls Mistral API with a list of prompts, concatenates results,
    formats the content, and generates a downloadable DOCX.
    """
    unique_key = f"{button_text.lower().replace(' ', '_')}_button"
    clicked = st.button(button_text, key=unique_key)

    # 🔍 Debug Info
    with st.expander("🧪 Debug Info (Prompt + State)"):
        st.write("🔘 **Button Clicked:**", "✅ Yes" if clicked else "❌ No")
        st.write("🙋 **User Confirmed Prompt:**", st.session_state.get("user_confirmation", False))
        st.write("🔑 **API Key Loaded:**", "✅ Yes" if api_key else "❌ No")

        for idx, prompt in enumerate(prompts):
            label = f"Prompt #{idx + 1}"
            st.write(f"📄 **{label} Exists:**", "✅ Yes" if prompt else "❌ No")
            if prompt:
                st.code(prompt[:500] + "..." if len(prompt) > 500 else prompt, language="markdown")

    if clicked:
        if not st.session_state.get("user_confirmation", False):
            st.warning("⚠️ Please confirm the AI prompt before generating the runbook.")
            return

        try:
            client = Mistral(api_key=api_key)
            combined_output = ""

            for idx, prompt in enumerate(prompts):
                st.info(f"📡 Querying Mistral for Section {idx + 1}...")
                response = client.chat.complete(
                    model="mistral-small-latest",
                    messages=[UserMessage(content=prompt)],
                    max_tokens=1500,
                    temperature=0.5,
                )
                combined_output += response.choices[0].message.content + "\n\n"

            formatted_output = format_output_for_docx(combined_output)

            st.success(f"{doc_heading} generated successfully!")
            st.write(combined_output)

            # Write to DOCX
            doc = Document()
            doc.add_heading(doc_heading, 0)
            doc.add_paragraph(formatted_output)
            doc.save(doc_filename)

            with open(doc_filename, "rb") as f:
                st.download_button(
                    label="📄 Download DOCX",
                    data=f,
                    file_name=doc_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"❌ Error generating runbook: {str(e)}")


#### Prompts Here #####

def query_utility_providers():
    """
    Queries Mistral AI for public utility providers based on city and ZIP code 
    stored in st.session_state. Stores and returns results in session state.
    """
    city = st.session_state.get("city", "").strip()
    zip_code = st.session_state.get("zip_code", "").strip()

    if not city or not zip_code:
        st.warning("City and ZIP code must be provided in session state.")
        return {
            "electricity": "Missing input",
            "natural_gas": "Missing input",
            "water": "Missing input"
        }

    prompt = f"""
You are a reliable assistant helping users prepare emergency documentation. 
Given the city: {city} and ZIP code: {zip_code}, list the **primary public utility provider companies** for the following:

1. Electricity
2. Natural Gas
3. Water

For each, provide only the company name. Format your response like this:

Electricity Provider: <company name>
Natural Gas Provider: <company name>
Water Provider: <company name>
"""

    try:
        response = client.chat.complete(
            model="mistral-small-latest",
            messages=[UserMessage(content=prompt)],
            max_tokens=1500,
            temperature=0.5,
        )
        content = response.choices[0].message.content
    except Exception as e:
        st.error(f"Error querying Mistral API: {str(e)}")
        content = ""

    def extract(label):
        match = re.search(rf"{label} Provider:\s*(.+)", content)
        return match.group(1).strip() if match else "Not found"

    electricity = extract("Electricity")
    natural_gas = extract("Natural Gas")
    water = extract("Water")

    st.session_state["electricity_provider"] = electricity
    st.session_state["natural_gas_provider"] = natural_gas
    st.session_state["water_provider"] = water

    return {
        "electricity": electricity,
        "natural_gas": natural_gas,
        "water": water
    }

def utilities_emergency_runbook_prompt(
            city=st.session_state.get("city", ""),
            zip_code=st.session_state.get("zip_code", ""),
            internet_provider_name=st.session_state.get("internet_provider",""),
            electricity_provider_name=st.session_state.get("electricity_provider",""),
            natural_gas_provider_name=st.session_state.get("natural_gas_provider",""),
            water_provider_name=st.session_state.get("water_provider","")
        ):

    return f"""
You are an expert assistant generating a city-specific Emergency Preparedness Run Book. First, search the internet for up-to-date local utility providers and their emergency contact information. Then, compose a comprehensive, easy-to-follow guide customized for residents of City: {city}, Zip Code: {zip_code}.

Start by identifying the following utility/service providers for the specified location:
- Internet Provider Name
- Electricity Provider Name
- Natural Gas Provider Name
- Water Provider Name

For each provider, retrieve:
- Company Description
- Customer Service Phone Number
- Customer Service Address (if available)
- Official Website
- Emergency Contact Numbers (specific to outages, leaks, service disruptions)
- Steps to report issues

---


### 📕 Emergency Run Book

#### ⚡ 1. Electricity – {electricity_provider_name}
- Provider Description
- Customer Service
- Website
- Emergency Contact

**Power Outage Response Guide:**
- Steps to follow
- How to report
- Safety precautions

---
#### 🔥 2. Natural Gas – {natural_gas_provider_name}
- Provider Description
- Customer Service
- Website
- Emergency Contact

**Gas Leak Response Guide:**
- Signs and precautions
- How to evacuate
- How to report

---
#### 💧 3. Water – {water_provider_name}
- Provider Description
- Customer Service
- Website
- Emergency Contact

**Water Outage or Leak Guide:**
- Detection steps
- Shutoff procedure

---
#### 🌐 4. Internet – {internet_provider_name}
- Provider Description
- Customer Service
- Website
- Emergency Contact

**Internet Outage Response Guide:**
- Troubleshooting
- Reporting
- Staying informed
---

Ensure the run book is clearly formatted using Markdown, with bold headers and bullet points. Use ⚠️ to highlight missing kit items.
""".strip()

#### Emergency Kit + Utilities Prompt ####

def emergency_kit_utilities_runbook_prompt(
    city=st.session_state.get("city", ""),
    zip_code=st.session_state.get("zip_code", ""),
    internet_provider_name=st.session_state.get("internet_provider",""),
    electricity_provider_name=st.session_state.get("electricity_provider",""),
    natural_gas_provider_name=st.session_state.get("natural_gas_provider",""),
    water_provider_name=st.session_state.get("water_provider",""),
    emergency_kit_status=st.session_state.get("emergency_kit_status", "No"),
    emergency_kit_location=st.session_state.get("emergency_kit_location", ""),
    selected_items=st.session_state.get("homeowner_kit_stock", []),
    not_selected_items=st.session_state.get("not_selected_items", []),
    additional_items=st.session_state.get("additional_kit_items", ""),
    flashlights_info=st.session_state.get("flashlights_info", ""),
    radio_info=st.session_state.get("radio_info", ""),
    food_water_info=st.session_state.get("food_water_info", ""),
    important_docs_info=st.session_state.get("important_docs_info", ""),
    whistle_info=st.session_state.get("whistle_info", ""),
    medications_info=st.session_state.get("medications_info", ""),
    mask_info=st.session_state.get("mask_info", ""),
    maps_contacts_info=st.session_state.get("maps_contacts_info", "")
):
    # Build Markdown lists for inventory & missing items
    selected_md = "".join(f"- {item}\n" for item in selected_items)
    missing_md  = "".join(f"- {item}\n" for item in not_selected_items)

    # Parse additional items (comma-separated) into bullets
    additional_list = [itm.strip() for itm in additional_items.split(",") if itm.strip()]
    additional_md   = "".join(f"- {itm}\n" for itm in additional_list)

    # Helper to render non-empty recommended items
    def render_recommended(*items):
        return "".join(f"- {i}\n" for i in items if i and i.strip())

    return f"""
You are an expert assistant generating a city-specific Emergency Preparedness Run Book. First, search the internet for up-to-date local utility providers and their emergency contact information. Then, compose a comprehensive, easy-to-follow guide customized for residents of City: {city}, Zip Code: {zip_code}.

Start by identifying the following utility/service providers for the specified location:
- Internet Provider Name
- Electricity Provider Name
- Natural Gas Provider Name
- Water Provider Name

For each provider, retrieve:
- Company Description
- Customer Service Phone Number
- Customer Service Address (if available)
- Official Website
- Emergency Contact Numbers (specific to outages, leaks, service disruptions)
- Steps to report issues

For Emergency Kit Summary:
{"Kit is available at " + emergency_kit_location if emergency_kit_status == "Yes" else "Kit is a work in progress and will be located at " + emergency_kit_location}

---

### 🧰 Emergency Kit Summary

**Kit Inventory:**  
{selected_md or "_(none selected)_"}  
⚠️ **Missing Kit Items (consider adding):**  
{missing_md or "_(none missing)_"}  

**Additional User-Added Items:**  
{additional_md or "_(none added)_"}  

---

### 📕 Emergency Run Book

#### ⚡ 1. Electricity – {electricity_provider_name}
- Provider Description
- Customer Service
- Website
- Emergency Contact

**Power Outage Response Guide:**
- Steps to follow
- How to report
- Safety precautions
- **Recommended Kit Items**:
{render_recommended(flashlights_info, radio_info, food_water_info, important_docs_info)}

---

#### 🔥 2. Natural Gas – {natural_gas_provider_name}
- Provider Description
- Customer Service
- Website
- Emergency Contact

**Gas Leak Response Guide:**
- Signs and precautions
- How to evacuate
- How to report
- **Recommended Kit Items**:
{render_recommended(whistle_info, important_docs_info, flashlights_info)}

---

#### 💧 3. Water – {water_provider_name}
- Provider Description
- Customer Service
- Website
- Emergency Contact

**Water Outage or Leak Guide:**
- Detection steps
- Shutoff procedure
- **Recommended Kit Items**:
{render_recommended(food_water_info, medications_info, mask_info, important_docs_info)}

---

#### 🌐 4. Internet – {internet_provider_name}
- Provider Description
- Customer Service
- Website
- Emergency Contact

**Internet Outage Response Guide:**
- Troubleshooting
- Reporting
- Staying informed
- **Recommended Kit Items**:
{render_recommended(radio_info, maps_contacts_info, important_docs_info)}

Ensure the run book is clearly formatted using Markdown, with bold headers and bullet points. Use ⚠️ to highlight missing kit items.
""".strip()


#### Mail + Trash Prompt ####
def mail_trash_runbook_prompt():
    mail_info = st.session_state.get("mail_info", {})
    mailbox_location = mail_info.get("Mailbox Location", "Not provided")
    mailbox_key = mail_info.get("Mailbox Key", "Not provided")
    pick_up_schedule = mail_info.get("Pick-Up Schedule", "Not provided")
    what_to_do_with_mail = mail_info.get("What to Do with the Mail", "Not provided")
    what_to_do_with_packages = mail_info.get("Packages", "Not provided")
    trash_info = st.session_state.get("trash_info", {})
    indoor = trash_info.get("indoor", {})
    outdoor = trash_info.get("outdoor", {})
    schedule = trash_info.get("schedule", {})
    composting = trash_info.get("composting", {})
    common_disposal = trash_info.get("common_disposal", {})
    wm = trash_info.get("waste_management", {})

    return f"""
You are an expert assistant generating Mail and Waste Management Run Book. Compose a comprehensive, easy-to-follow guide for house stitters and people watching the house when occupants are out of town. For any values set to No please omit thoses lines.

### 📕 Mail Handling and Waste Management Instructions 

#### 📬 Mail Handling Instructions

- **Mailbox Location**: {mailbox_location}
- **Mailbox Key Info**: {mailbox_key}
- **Pick-Up Schedule**: {pick_up_schedule}
- **Mail Sorting Instructions**: {what_to_do_with_mail}
- **Delivery Packages**: {what_to_do_with_packages}

---

#### 🗑️ Trash & Recycling Instructions

**Indoor Trash**
- Kitchen Trash: {indoor.get("kitchen_bin", "Not provided")}
- Bathroom Trash: {indoor.get("bathroom_bin", "Not provided")}
- Other Rooms Trash: {indoor.get("other_room_bin", "Not provided")}

**Outdoor Bins**
- Please take the bins: {outdoor.get("bin_destination", "Not provided")}
- Bins Description: {outdoor.get("bin_description", "Not provided")}
- Location: {outdoor.get("bin_location_specifics", "Not provided")}
- Instructions: {outdoor.get("bin_handling_instructions", "Not provided")}

**Collection Schedule**
- Garbage Pickup: {schedule.get("trash_day", "Not provided")}, {schedule.get("trash_time", "Not provided")}
- Recycling Pickup: {schedule.get("recycling_day", "Not provided")}, {schedule.get("recycling_time", "Not provided")}

**Composting**
- Composting Used: {"Yes" if composting.get("compost_used", False) else "No"}
- Compost Instructions: {composting.get("compost_instructions", "N/A")}

**Common Disposal Area**
- Used: {"Yes" if common_disposal.get("uses_common_disposal", False) else "No"}
- Instructions: {common_disposal.get("common_area_instructions", "N/A")}

**Waste Management Contact**
- Company Name: {wm.get("company_name", "Not provided")}
- Phone: {wm.get("phone", "Not provided")}
- Contact: {wm.get("description", "Not provided")}

---

Ensure the run book is clearly formatted using Markdown, with bold headers and bullet points. Use ⚠️ to highlight missing kit items.
""".strip()

#### Security and Services Prompt ####

def home_caretaker_runbook_prompt():
    csi = st.session_state.get("convenience_seeker_info", {})
    roi = st.session_state.get("rent_own_info", {})
    hsi = st.session_state.get("home_security_info", {})

    return f"""
You are a helpful assistant tasked with generating a professional, detailed, and easy-to-follow Home Caretaker & Guest Runbook. The goal is to ensure a smooth experience for caretakers or guests while the home occupants are away. 

Please use the following information provided by the homeowner to write a clear and structured guide:
Please omit any headings that return "Not provided" for all the values below it.
Please omit any sub-headings that return "Not provided" for all the values below it.
Please omit any lines that return "Not provided" or "N/A".
Please omit any sub-headings that return "Not provided" or "N/A" for all the values below it.
Please don't add a title to the runbook.

### 📕 Security and Services Guide

#### 🔐 Home Security & Technology
- Security Company Name: {hsi.get("home_security_comp_name", "Not provided")}
- Security Company Number: {hsi.get("home_security_comp_num", "Not provided")}
- Arming/Disarming Instructions: {hsi.get("arm_disarm_instructions", "Not provided")}
- If Alert is Triggered: {hsi.get("security_alert_steps", "Not provided")}
- Indoor Camera Notes: {hsi.get("indoor_cameras", "Not provided")}
- Emergency Access Instructions: {hsi.get("access_emergency", "Not provided")}
- Wi-Fi Info Location: {hsi.get("wifi_network_name", "Not provided")}
- Guest Wi-Fi Access: {hsi.get("wifi_guests", "Not provided")}
- Landline/VOIP Notes: {hsi.get("landline_voip", "Not provided")}

---

#### 🧹 Cleaning Service Instructions
- Company Name: {csi.get("cleaning_name", "Not provided")}
- Phone Number: {csi.get("cleaning_number", "Not provided")}
- Schedule: {csi.get("cleaning_schedule", "Not provided")}
- Access Method: {csi.get("cleaning_access", "Not provided")}
- Post-Cleaning Procedures: {csi.get("cleaning_finish_steps", "Not provided")}
- Crew Identity Verification: {csi.get("cleaning_identity_confirmation", "Not provided")}

---

#### 🌿 Gardening & Landscape Service Instructions
- Company Name: {csi.get("gardening_name", "Not provided")}
- Phone Number: {csi.get("gardening_number", "Not provided")}
- Schedule: {csi.get("gardening_schedule", "Not provided")}
- Access Method: {csi.get("gardening_access", "Not provided")}
- Post-Service Procedures: {csi.get("gardening_finish_steps", "Not provided")}
- Crew Identity Verification: {csi.get("gardening_identity_confirmation", "Not provided")}

---

#### 🏊 Pool Maintenance Instructions
- Company Name: {csi.get("pool_name", "Not provided")}
- Phone Number: {csi.get("pool_number", "Not provided")}
- Schedule: {csi.get("pool_schedule", "Not provided")}
- Access Method: {csi.get("pool_access", "Not provided")}
- Post-Service Procedures: {csi.get("pool_finish_steps", "Not provided")}
- Crew Identity Verification: {csi.get("pool_identity_confirmation", "Not provided")}

---

#### 🏢 Property Management (Renters or HOA)
- Company Name: {roi.get("property_management_name", "Not provided")}
- Phone Number: {roi.get("property_management_number", "Not provided")}
- Email: {roi.get("property_management_email", "Not provided")}
- When to Contact: {roi.get("property_management_description", "Not provided")}

---

#### 🛠️ Service Contacts (For Homeowners)
**Handyman**
- Name: {roi.get("handyman_name", "N/A")}
- Phone: {roi.get("handyman_number", "N/A")}
- When to Contact: {roi.get("handyman_description", "N/A")}

**Electrician**
- Name: {roi.get("electrician_name", "N/A")}
- Phone: {roi.get("electrician_number", "N/A")}
- When to Contact: {roi.get("electrician_description", "N/A")}

**Exterminator**
- Name: {roi.get("exterminator_name", "N/A")}
- Phone: {roi.get("exterminator_number", "N/A")}
- When to Contact: {roi.get("exterminator_description", "N/A")}

**Plumber**
- Name: {roi.get("plumber_name", "N/A")}
- Phone: {roi.get("plumber_number", "N/A")}
- When to Contact: {roi.get("plumber_description", "N/A")}

---

Please format the runbook clearly with headers and bullet points. Use “⚠️ Not provided” as a flag for incomplete or missing info that should be reviewed.
""".strip()

#### Emergency Kit Documents ####

def emergency_kit_document_prompt():
    """
    Build the LLM prompt string from session_state, but first prune
    document_details so that each document only includes the fields
    for the storage locations it’s actually assigned to.
    """
    intro = (
        "Welcome to your Emergency Document Kit.\n\n"
        "- Reduces stress, and delivers peace of mind in a crisis.\n\n"
        "How to read this document:\n"
        "1. Start here for the value of the kit:\n\n"
        "   • Provides clear, step-by-step guidance to locate and retrieve vital documents quickly in a crisis.\n"
        "2. Scroll down to each storage location heading (e.g., “## Safe”)—"
        "these are sorted by where most documents live first.\n"
        "3. Under each location, you’ll find:\n\n"
        "   • Location details (address or placement info)\n\n"
        "   • Platform (if applicable, e.g., cloud service or password manager)\n\n"
        "   • Access steps (what’s required in an emergency)\n\n"
        "   • Contents notes (if there are multiple containers)\n\n"
        "4. The final list shows only the documents stored there, with categories.\n"
        "Keep this kit handy and review periodically to ensure accuracy."
    )

    global_physical  = st.session_state.get("global_physical_storage", [])
    global_digital   = st.session_state.get("global_digital_storage", [])
    raw_details      = st.session_state.get("document_details", {})

    # Build a filtered version
    filtered_details = {}
    for doc, details in raw_details.items():
        assigned = details.get("assigned_storage", [])
        fd = {"assigned_storage": assigned}

        # for every other key in details, only keep it if:
        # 1) it’s non-empty, and
        # 2) its key starts with one of the assigned-location prefixes
        for key, val in details.items():
            if key == "assigned_storage" or not val:
                continue
            for loc in assigned:
                prefix = loc.lower().replace(" ", "_").replace("/", "_") + "_"
                if key.startswith(prefix):
                    fd[key] = val
                    break

        filtered_details[doc] = fd

    phys_json    = json.dumps(global_physical, indent=2)
    digi_json    = json.dumps(global_digital, indent=2)
    details_json = json.dumps(filtered_details,  indent=2)

    return f"""
You are an expert at creating clear, action-ready Emergency Document Kits.

Below is the **introductory section** you must include exactly as written at the top of your output (no backticks):

{intro}

You will then be provided with three Python variables in JSON form:

global_physical_storage = {phys_json}  
global_digital_storage  = {digi_json}  
document_details        = {details_json}  

> **Important:** Only use the `document_details` mapping—do **not** pull from any other list.

**Your task**  
1. Output the introductory section verbatim as the first lines.  
2. Group all documents by storage location, showing physical first, then digital.  
3. Within each location, list only the documents actually stored there and include **only** these subsections **when they have data**:
   - **Location details:** the user-provided address or placement info  
   - **Platform:** the service or tool used (if present)  
   - **Access steps:** the emergency steps or authorizations required  
   - **Contents:** if multiple, what each container holds  
4. Sort locations by document count (most → least).  
5. Skip any location with zero documents.  
6. Format the rest in plain Markdown (no code fences), with one top-level heading per location:

   ## <Storage Name> (n documents)

   **Location details:**  
   <location or branch_address>  

   **Platform:**  
   <platform name>  

   **Access steps:**  
   <access_steps>  

   **Contents:**  
   <contents note>  

   **Documents stored:**  
   - **<Document A>** (Category: <Category>)  
   - **<Document B>** (Category: <Category>)

*Omit any subsection line if that field is empty or not applicable.*

Begin your response now.
""".strip()

def bonus_level_runbook_prompt():
    bonus_info = st.session_state.get("bonus_info", {})

    # Filter out empty entries
    filtered_info = {
        key: val
        for key, val in bonus_info.items()
        if val and str(val).strip()
    }
    bonus_json = json.dumps(filtered_info, indent=2)

    prompt = f"""
Append the following **Additional Information** to the home caretaker emergency runbook.  
Use clear headings and actionable bullet points.  
**Only include a section if its data appears in the JSON. Omit any heading with no data.**

Additional Information (JSON):
{bonus_json}

Your task:
1. **Home Maintenance**  
   - Only if `maintenance_tasks` or `appliance_instructions` exist, under a **Home Maintenance** heading list those items.
2. **Home Rules & Preferences**  
   - Only if `house_rules` or `cultural_practices` exist, under a **Home Rules & Preferences** heading summarize them.
3. **Housekeeping & Cleaning**  
   - Only if `housekeeping_instructions` or `cleaning_preferences` exist, under a **Housekeeping & Cleaning** heading provide routines and supply locations.
4. **Entertainment & Technology**  
   - Only if `entertainment_info` or `device_instructions` exist, under an **Entertainment & Technology** heading describe usage and charging steps.

Produce **only** the formatted runbook addition, starting with:

## Additional Information

…followed by the pertinent sections (no blank headings, no extra commentary).
""".strip()

    return prompt

###### Main Functions that comprise of the Levels

### Leve 1 - Home

def home_debug():

    st.write("🟡 About to render runbook button")

    generate_runbook_from_prompt(
        prompt=st.session_state.get("generated_prompt", ""),
        api_key=os.getenv("MISTRAL_TOKEN"),
        button_text="Complete Level 1 Mission",
        doc_heading="Home Utilities Emergency Runbook",
        doc_filename="home_utilities_emergency.html"
    )
    st.write("🟢 After button render")

def home():
    st.write("Let's gather some information. Please enter your details:")

    # Input fields
    st.session_state.city = st.text_input("City", value=st.session_state.get("city", ""))
    st.session_state.zip_code = st.text_input("ZIP Code", value=st.session_state.get("zip_code", ""))
    st.session_state.internet_provider = st.text_input("Internet Provider", value=st.session_state.get("internet_provider", ""))

    # Step 1: Fetch utility providers
    if st.button("Find My Utility Providers"):
        with st.spinner("Fetching providers from Mistral..."):
            results = query_utility_providers()
            st.success("Providers stored in session state!")

    # Step 2: Allow corrections
    st.write("Correct Utility Providers:")

    correct_electricity = st.checkbox("Correct Electricity Provider", value=False)
    corrected_electricity = st.text_input("Electricity Provider", value=st.session_state.get("electricity_provider", ""), disabled=not correct_electricity)

    correct_natural_gas = st.checkbox("Correct Natural Gas Provider", value=False)
    corrected_natural_gas = st.text_input("Natural Gas Provider", value=st.session_state.get("natural_gas_provider", ""), disabled=not correct_natural_gas)

    correct_water = st.checkbox("Correct Water Provider", value=False)
    corrected_water = st.text_input("Water Provider", value=st.session_state.get("water_provider", ""), disabled=not correct_water)

    if st.button("Save Utility Providers"):
        if correct_electricity:
            st.session_state["electricity_provider"] = corrected_electricity
        if correct_natural_gas:
            st.session_state["natural_gas_provider"] = corrected_natural_gas
        if correct_water:
            st.session_state["water_provider"] = corrected_water
        st.success("Utility providers updated!")

    # Step 3: Preview prompt
    # Move this outside the expander
    user_confirmation = st.checkbox("✅ Confirm AI Prompt")
    st.session_state["user_confirmation"] = user_confirmation # store confirmation in session

    if user_confirmation:
        prompt = utilities_emergency_runbook_prompt()
        st.session_state["generated_prompt"] = prompt
    else:
        st.session_state["generated_prompt"] = None

    st.session_state.progress["level_1_completed"] = True
    save_progress(st.session_state.progress)


# Show prompt in expander
    with st.expander("AI Prompt Preview (Optional)"):
        if st.session_state.get("generated_prompt"):
            st.code(st.session_state["generated_prompt"], language="markdown")

    # Step 4: Generate runbook using reusable function
    st.write("Next, click the button to generate your personalized utilities emergency runbook document:")
    
    if not st.session_state.get("generated_prompt"):
        st.warning("⚠️ Prompt not ready. Please confirm the prompt first.")
        return
    
    #st.write("Prompt preview (sanity check):", st.session_state.get("generated_prompt", "[Empty]"))

    generate_runbook_from_prompt(
        prompt=st.session_state.get("generated_prompt", ""),
        api_key=os.getenv("MISTRAL_TOKEN"),
        button_text="Complete Level 1 Mission",
        doc_heading="Home Utilities Emergency Runbook",
        doc_filename="home_utilities_emergency.docx"
    )
    #st.write("🧪 Debug Info:")
    #st.write("Prompt exists:", "Yes" if st.session_state.get("generated_prompt") else "No")
    #st.write("User confirmed:", st.session_state.get("user_confirmation"))
    #st.write("Prompt:", st.session_state.get("generated_prompt"))
    #st.write("API key loaded:", "Yes" if os.getenv("MISTRAL_TOKEN") else "No")


### Level 2 - Emergency Kit Details

# Define the homeowner_kit_stock function
import streamlit as st

def homeowner_kit_stock():
    kit_items = [
        "Flashlights and extra batteries",
        "First aid kit",
        "Non-perishable food and bottled water",
        "Medications and personal hygiene items",
        "Important documents (insurance, identification)",
        "Battery-powered or hand-crank radio",
        "Whistle (for signaling)",
        "Dust masks (for air filtration)",
        "Local maps and contact lists"
    ]

    # Initialize storage keys in session_state
    for item in kit_items:
        storage_key = (
            item.lower()
                .replace(" ", "_")
                .replace("(", "")
                .replace(")", "")
            + "_storage"
        )
        if storage_key not in st.session_state:
            st.session_state[storage_key] = None

    with st.form(key="emergency_kit_form"):
        st.write("Select all you have:")
        selected = []

        # Display checkboxes in rows of 4
        for start in range(0, len(kit_items), 4):
            chunk = kit_items[start : start + 4]
            cols = st.columns(len(chunk))
            for idx, item in enumerate(chunk):
                storage_key = (
                    item.lower()
                        .replace(" ", "_")
                        .replace("(", "")
                        .replace(")", "")
                    + "_storage"
                )
                # default checked if previously stored
                default_checked = st.session_state[storage_key] is not None
                checked = cols[idx].checkbox(
                    item,
                    value=default_checked,
                    key=f"chk_{storage_key}"
                )
                if checked:
                    selected.append(item)

        submit = st.form_submit_button(label="Submit")

    if submit:
        not_selected = [item for item in kit_items if item not in selected]
        if not_selected:
            st.warning("⚠️ Consider adding the following items to your emergency kit:")
            for item in not_selected:
                st.write(f"- {item}")

        # Update session_state based on checkboxes
        for item in kit_items:
            storage_key = (
                item.lower()
                    .replace(" ", "_")
                    .replace("(", "")
                    .replace(")", "")
                + "_storage"
            )
            if item in selected:
                st.session_state[storage_key] = item
            else:
                st.session_state[storage_key] = None

    return selected

def emergency_kit():
    # Use st.radio to create a dropdown menu for selecting between renting or owning
    emergency_kit_status = st.radio(
        'Do you have an Emergency Kit?',
        ('Yes', 'No')
    )

    kit_items = [
        "Flashlights and extra batteries",
        "First aid kit",
        "Non-perishable food and bottled water",
        "Medications and personal hygiene items",
        "Important documents (insurance, identification)",
        "Battery-powered or hand-crank radio",
        "Whistle (for signaling)",
        "Dust masks (for air filtration)",
        "Local maps and contact lists"
    ]

    st.write("🧰 Emergency Kit Info")
    if emergency_kit_status == 'Yes':
        st.success('Great—you already have a kit!', icon=":material/medical_services:")
    else:
        st.warning("⚠️ Let's build your emergency kit with what you have.")

    st.session_state['emergency_kit_status'] = emergency_kit_status

    # Kit location
    emergency_kit_location = st.text_area("Where is (or where will) the Emergency Kit be located?")
    if emergency_kit_location:
        st.session_state['emergency_kit_location'] = emergency_kit_location

    # Core kit items selector
    selected_items = homeowner_kit_stock()
    if selected_items is not None:
        st.session_state['homeowner_kit_stock'] = selected_items

    # Additional custom items
    additional = st.text_input(
        "Add any additional emergency kit items not in the list above (comma-separated):",
        value=st.session_state.get('additional_kit_items', '')
    )
    if additional is not None:
        # store raw string or split into list
        st.session_state['additional_kit_items'] = additional

    # Compute not selected
    current_selection = st.session_state.get('homeowner_kit_stock', [])
    not_selected_items = [item for item in kit_items if item not in current_selection]
    st.session_state['not_selected_items'] = not_selected_items

    return not_selected_items

def emergency_kit_utilities():

    # Step 1: Input fields
    emergency_kit()
    
    # Step 2: Preview prompt

    # Move this outside the expander
    user_confirmation = st.checkbox("✅ Confirm AI Prompt")
    st.session_state["user_confirmation"] = user_confirmation # store confirmation in session

    st.session_state.progress["level_2_completed"] = True
    save_progress(st.session_state.progress)

    if user_confirmation:
        prompt = emergency_kit_utilities_runbook_prompt()
        st.session_state["generated_prompt"] = prompt
    else:
        st.session_state["generated_prompt"] = None

# Show prompt in expander
    with st.expander("AI Prompt Preview (Optional)"):
        if st.session_state.get("generated_prompt"):
            st.code(st.session_state["generated_prompt"], language="markdown")

    # Step 3: Generate runbook using reusable function
    st.write("Next, click the button to generate your personalized utilities emergency runbook document:")

    generate_runbook_from_prompt(
        prompt=st.session_state.get("generated_prompt", ""),
        api_key=os.getenv("MISTRAL_TOKEN"),
        button_text="Complete Level 2 Mission",
        doc_heading="Home Emergency Runbook With Emergency Kit Summary",
        doc_filename="home_util_emergency_kit.docx"
    )
##### Level 3 - Mail Handling and Trash

def mail():

    if 'mail_info' not in st.session_state:
        st.session_state.mail_info = {}

    with st.expander("Mail Handling", expanded=True):
        # Input fields
        mailbox_location = st.text_area(
            "📍 Mailbox Location",
            placeholder="E.g., 'At the end of the driveway on the left side.'"
        )

        mailbox_key = st.text_area(
            "🔑 Mailbox Key (Optional)",
            placeholder="E.g., 'Hanging on the key hook next to the fridge.'"
        )

        pick_up_schedule = st.text_area(
            "📆 Mail Pick-Up Schedule",
            placeholder="E.g., 'Every other day' or 'Mondays and Thursdays'"
        )

        what_to_do_with_mail = st.text_area(
            "📥 What to Do with the Mail",
            placeholder="E.g., 'Place it in the tray on the kitchen counter.'"
        )

        What_to_do_with_packages = st.text_area(
            "📦 Packages",
            placeholder="E.g., 'Place it inside the entryway closet.'"
        )

        # Dynamic progress bar based on completion
        completed = sum([
            bool(mailbox_location),
            bool(mailbox_key),
            bool(pick_up_schedule),
            bool(what_to_do_with_mail),
            bool(What_to_do_with_packages)
        ])
        progress = int((completed / 5) * 100)
        st.progress(progress)

        # Save button
        if st.button("✅ Mail Handling 100% Complete. Click to Save"):
            st.session_state.mail_info = {
                "Mailbox Location": mailbox_location,
                "Mailbox Key": mailbox_key,
                "Pick-Up Schedule": pick_up_schedule,
                "What to Do with the Mail": what_to_do_with_mail,
                "Packages": What_to_do_with_packages
            }
            st.success("Mail handling instructions saved successfully!")

    # Display saved info
    st.subheader("📂 Saved Mail Handling Information")
    if st.session_state.mail_info:
        with st.expander("📋 Review Saved Info", expanded=True):
            for key, value in st.session_state.mail_info.items():
                st.markdown(f"**{key}:** {value}")
    else:
        st.info("No mail handling information saved yet.")

def trash_handling():

    if 'trash_info' not in st.session_state:
        st.session_state.trash_info = {}
    if 'trash_images' not in st.session_state:
        st.session_state.trash_images = {}

    days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
    times = ["Morning", "Afternoon", "Evening"]

    # --- Indoor Trash Info ---
    with st.expander("Kitchen and Bath Trash Details", expanded=True):
        st.markdown("##### Fill in the kitchen and bathroom trash info")

        kitchen_bin = st.text_area(
            "Kitchen Trash Bin Location, Emptying Schedule and Replacement Trash Bags", 
            placeholder="E.g. Bin is located under the kitchen sink. Empty when full.  Bags are next to the bin. They are labeled kitchen bags."
        )

        bathroom_bin = st.text_area(
            "Bathroom Trash Bin Emptying Schedule and Replacement Trash Bags ", 
            placeholder="E.g. Empty before Trash day.  Bags are under the sink."
        )

        other_room_bin = st.text_area(
            "Other Room Trash Bin Emptying Schedule and Replacement Trash Bags ", 
            placeholder="E.g. Empty before Trash day.  Bags are under the sinks of each bathroom."
        )

    # --- Outdoor Bin Info ---
    with st.expander("Outdoor Bin Details", expanded=True):
        st.markdown("##### Outdoor Bin Handling Details")

        bin_description = st.text_area(
            "What the Outdoor Trash Bins Look Like", 
            placeholder="E.g. Green with lid"
        )

        bin_location_specifics = st.text_area(
            "Specific Location or Instructions for Outdoor Bins", 
            placeholder="E.g. Next to side gate"
        )

        # Image upload helper
        def handle_image(label, display_name):
            image_key = f"{label} Image"
            if image_key not in st.session_state.trash_images:
                st.session_state.trash_images[image_key] = None

            if st.session_state.trash_images[image_key]:
                st.image(Image.open(io.BytesIO(st.session_state.trash_images[image_key])), caption=display_name)
                if st.button(f"Delete {display_name}", key=f"delete_{label}"):
                    st.session_state.trash_images[image_key] = None
                    st.experimental_rerun()
            else:
                uploaded = st.file_uploader(f"Upload a photo of the {display_name}", type=["jpg", "jpeg", "png"], key=f"{label}_upload")
                if uploaded:
                    st.session_state.trash_images[image_key] = uploaded.read()
                    st.success(f"{display_name} image uploaded.")
                    st.experimental_rerun()

        handle_image("Outdoor Bin", "Outdoor Trash Bin")
        handle_image("Recycling Bin", "Recycling Bin")

    # --- Collection Schedule ---
    with st.expander("Collection Schedule", expanded=True):
        st.markdown("##### Enter your trash and recycling schedule")
       
       #Collection Day + Time Pickers
        trash_day = st.selectbox("Garbage Pickup Day", days)
        trash_time = st.selectbox("Garbage Pickup Time", times)
        recycling_day = st.selectbox("Recycling Pickup Day", days)
        recycling_time = st.selectbox("Recycling Pickup Time", times)

        bin_handling_instructions = st.text_area("Instructions for Placing and Returning Outdoor Bins")

    # --- Common Disposal Area ---
    with st.expander("Common Disposal Area (if applicable)", expanded=True):
        st.markdown("##### Shared disposal area details")
        uses_common_disposal = st.checkbox("Is there a common disposal area?")

        if uses_common_disposal:
            common_area_instructions = st.text_area(
                "Instructions for Common Disposal Area", 
                placeholder="E.g. Dumpster in alley"
            )
            if common_area_instructions:
                handle_image("Common Area", "Common Disposal Area")

    # --- Composting ---
    with st.expander("Composting Instructions (if applicable)", expanded=True):
        st.markdown("##### Composting info")
        compost_applicable = st.checkbox("Is composting used?")

        if compost_applicable:
            compost_instructions = st.text_area(
                "Compost Instructions", 
                placeholder="E.g. Put organics in green bin"
            )

    # --- Waste Management Contact ---
    with st.expander("Waste Management Contact Info", expanded=True):
        st.markdown("##### Company contact details")

        wm_name = st.text_input(
            "Waste Management Company Name", 
            placeholder="E.g. WastePro"
        )
        wm_phone = st.text_input(
            "Contact Phone Number", 
            placeholder="E.g. (123) 456-7890"
            )
        wm_description = st.text_area(
            "When to Contact", 
            placeholder="E.g. Missed pickup or billing issues"
            )

# Dynamic progress bar based on completion
        total_sections = 6  # or 7 if compost is enabled
        filled_sections = sum([
            bool(kitchen_bin),
            bool(bathroom_bin),
            bool(other_room_bin),
            bool(bin_description),
            bool(bin_location_specifics),
            bool(trash_day),
            bool(trash_time),
            bool(recycling_day),
            bool(recycling_time),
            bool(wm_name),
            bool(wm_phone),
            bool(wm_description),
            compost_applicable and bool(compost_instructions),
            uses_common_disposal and bool(common_area_instructions),
        ])

        total_progress = int((filled_sections / 16) * 100)
        st.progress(total_progress)

    # --- Save Button ---
    if st.button("✅ Trash Handling 100% Complete. Click to Save"):
        st.session_state.trash_info = {
            "indoor": {
                "kitchen_bin": kitchen_bin,
                "bathroom_bin": bathroom_bin,
                "other_room_bin": other_room_bin,
            },
            "outdoor": {
                "bin_description": bin_description,
                "bin_location_specifics": bin_location_specifics,
                "bin_handling_instructions": bin_handling_instructions
            },
            "schedule": {
                "trash_day": trash_day,
                "trash_time": trash_time,
                "recycling_day": recycling_day,
                "recycling_time": recycling_time
            },
            "composting": {
                "compost_used": compost_applicable,
                "compost_instructions": compost_instructions if compost_applicable else "N/A"
            },
            "common_disposal": {
                "uses_common_disposal": uses_common_disposal,
                "common_area_instructions": common_area_instructions if uses_common_disposal else "N/A"
            },
            "waste_management": {
                "company_name": wm_name,
                "phone": wm_phone,
                "description": wm_description
            }
    }

    st.session_state.progress["trash_completed"] = True
    save_progress(st.session_state.progress)

    # --- Display saved info and uploaded images ---
    if st.session_state.trash_info:
        st.markdown("### ✅ Saved Trash Handling Information")
        for key, value in st.session_state.trash_info.items():
            st.write(f"**{key}**: {value}")

    if st.session_state.trash_images:
        st.write("🖼️ Uploaded Photos")
        for label, image_bytes in st.session_state.trash_images.items():
            if image_bytes:
                st.image(Image.open(io.BytesIO(image_bytes)), caption=label)

def mail_trash_handling():
    st.write("📬 Level 3: Mail & Trash Handling")

    # Ensure session_state keys exist
    st.session_state.progress.setdefault("level_3_completed", False)
    st.session_state.setdefault("user_confirmation", False)
    st.session_state.setdefault("prompt_emergency", None)
    st.session_state.setdefault("prompt_mail_trash", None)

    # Create three tabs
    tab1, tab2, tab3 = st.tabs([
        "📬 Mail Input",
        "🗑️ Trash Input",
        "🤖 Review & Generate"
    ])

    # ─── Tab 1: Mail Input ─────────────────────────────────────
    with tab1:
        st.subheader("Step 1: 📬 Mail Handling Instructions")
        mail()  # your existing mail() form

    # ─── Tab 2: Trash Input ────────────────────────────────────
    with tab2:
        st.subheader("Step 2: 🗑️ Trash Disposal Instructions")
        trash_handling()  # your existing trash_handling() form

    # ─── Tab 3: Prompt Review & Generate ───────────────────────
    with tab3:
        st.subheader("Step 3: Review & Generate Runbook")

        # Confirm before generating prompts
        confirmed = st.checkbox(
            "✅ Confirm AI Prompt",
            value=st.session_state.user_confirmation,
            key="mail_trash_confirm"
        )
        st.session_state.user_confirmation = confirmed

        if confirmed:
            # mark level complete & build prompts
            st.session_state.progress["level_3_completed"] = True
            save_progress(st.session_state.progress)
            st.session_state.prompt_emergency = emergency_kit_utilities_runbook_prompt()
            st.session_state.prompt_mail_trash = mail_trash_runbook_prompt()

            # Show prompt preview
            st.markdown("#### 🆘 Emergency + Utilities Prompt")
            st.code(st.session_state.prompt_emergency, language="markdown")

            st.markdown("#### 📬 Mail + Trash Prompt")
            st.code(st.session_state.prompt_mail_trash, language="markdown")

            # Generate runbook button
            generate_runbook_from_prompt_split(
                prompt_emergency=st.session_state.prompt_emergency,
                prompt_mail_trash=st.session_state.prompt_mail_trash,
                api_key=os.getenv("MISTRAL_TOKEN"),
                button_text="Complete Level 3 Mission",
                doc_heading="Home Emergency Runbook for Caretakers and Guests",
                doc_filename="home_runbook_caretakers.docx"
            )
        else:
            st.info("Please confirm the AI prompt after entering your details in Tabs 1 and 2.")

##### Level 4 - Home Security and Services

def home_security():
    st.write("💝 Security-Conscious")

    # Initialize session state
    if 'home_security_info' not in st.session_state:
        st.session_state.home_security_info = {}

    with st.expander("Home Security System (if applicable)", expanded=True):
        st.markdown("##### Home Security and Privacy Info")
        home_security_applicable = st.checkbox("Are you home security and privacy conscious?")

        if home_security_applicable:
            info = st.session_state.home_security_info
            info['home_security_applicable'] = True

            # ─── Company Details ────────────────────────────────────
            st.subheader("🏢 Company Details")
            info['home_security_comp_name'] = st.text_input(
                "Security Company Name",
                value=info.get('home_security_comp_name', '')
            )
            info['home_security_comp_num'] = st.text_input(
                "Security Company Phone Number",
                value=info.get('home_security_comp_num', '')
            )

            # ─── System Operation ───────────────────────────────────
            st.subheader("⚙️ System Operation")
            info['arm_disarm_instructions'] = st.text_area(
                "Instructions to arm/disarm system",
                placeholder="e.g., via app, keypad code, shared link",
                value=info.get('arm_disarm_instructions', '')
            )
            info['security_alert_steps'] = st.text_area(
                "Steps if a security alert is triggered",
                placeholder="e.g., check app, contact company",
                value=info.get('security_alert_steps', '')
            )
            info['indoor_cameras'] = st.text_area(
                "Indoor cameras/monitoring details and activation",
                placeholder="e.g., motion sensors, smartphone access",
                value=info.get('indoor_cameras', '')
            )

            # ─── Emergency Access ──────────────────────────────────
            st.subheader("🚨 Emergency & Lockout Access")
            info['access_emergency'] = st.text_area(
                "Emergency access instructions & storage location",
                placeholder="e.g., spare key in lockbox, PIN in password manager",
                value=info.get('access_emergency', '')
            )

            # ─── Network Information ───────────────────────────────
            st.subheader("📶 Network Information")
            info['wifi_network_location'] = st.text_input(
                "Where is Wi-Fi network name/password stored?",
                value=info.get('wifi_network_location', '')
            )
            info['wifi_guests'] = st.text_input(
                "Guest network details & password sharing method",
                value=info.get('wifi_guests', '')
            )

            # ─── Phone Setup ───────────────────────────────────────
            st.subheader("📞 Home Phone / VoIP")
            info['landline_voip'] = st.text_area(
                "Home phone setup & call-handling instructions",
                placeholder="e.g., handsets, contact for issues",
                value=info.get('landline_voip', '')
            )

            # ─── Save Button ───────────────────────────────────────
            if st.button("💾 Save Home Security Info"):
                st.success("✅ Home security information saved!")
        else:
            st.info("🔒 You indicated home security is not applicable.")
            st.session_state.home_security_info = {"home_security_applicable": False}

def convenience_seeker():
    st.write("🧼 Quality-Oriented Household Services")

    # Initialize in session state
    if 'convenience_seeker_info' not in st.session_state:
        st.session_state.convenience_seeker_info = {}
    if 'convenience_seeker_options' not in st.session_state:
        st.session_state.convenience_seeker_options = []

    with st.expander("Home Quality-Oriented (if applicable)", expanded=True):
        st.markdown("##### Services You Invest In")
        services = ["Cleaning", "Gardening/Landscape", "Pool Maintenance"]

        # Multi-select segmented control
        selected_services = st.segmented_control(
            "As someone who wants their home and garden to be well‐maintained "
            "and is willing to invest in professional help, what services do you pay for?",
            options=services,
            selection_mode="multi",
            default=st.session_state.convenience_seeker_options,
            key="convenience_seeker_options"
        )
        # Save selection
        st.session_state.convenience_seeker_info['convenience_seeker_options'] = selected_services

        # --- Cleaning Service ---
        if "Cleaning" in selected_services:
            st.subheader("🧹 Cleaning Service Info")
            info = st.session_state.convenience_seeker_info

            info['cleaning_name'] = st.text_input(
                "Cleaning Company Name",
                value=info.get('cleaning_name', '')
            )

            info['cleaning_number'] = st.text_input(
                "Cleaning Company Phone Number",
                value=info.get('cleaning_number', '')
            )

            # Cleaning Frequency
            freq_options = ["Monthly", "Bi-Weekly", "Weekly"]
            default_freq = info.get('cleaning_frequency', freq_options[0])
            info['cleaning_frequency'] = st.selectbox(
                "Cleaning Frequency",
                options=freq_options,
                index=freq_options.index(default_freq),
                key="cleaning_frequency"
            )

            # Day of the Week
            days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday", "Not Specified"]
            default_day = info.get('cleaning_day_of_week', days[0])
            info['cleaning_day_of_week'] = st.selectbox(
                "Cleaning Day of the Week",
                options=days,
                index=days.index(default_day),
                key="cleaning_day_of_week"
            )

            info['cleaning_access'] = st.text_input(
                "Access Method for Cleaners",
                value=info.get('cleaning_access', '')
            )
            info['cleaning_finish_steps'] = st.text_area(
                "Post-Cleaning Procedures",
                value=info.get('cleaning_finish_steps', '')
            )
            info['cleaning_identity_confirmation'] = st.text_area(
                "Cleaning Crew Identity Verification",
                value=info.get('cleaning_identity_confirmation', '')
            )

        # --- Gardening/Landscape Service ---
        if "Gardening/Landscape" in selected_services:
            st.subheader("🌿 Gardening/Landscape Service Info")
            info = st.session_state.convenience_seeker_info

            info['gardening_name'] = st.text_input(
                "Gardening Company Name",
                value=info.get('gardening_name', '')
            )
            info['gardening_number'] = st.text_input(
                "Gardening Company Phone Number",
                value=info.get('gardening_number', '')
            )

            # → New: Gardening Frequency
            freq_options = ["Monthly", "Bi-Weekly", "Weekly"]
            default_freq = info.get('gardening_frequency', freq_options[0])
            frequency = st.selectbox(
                "Gardening Frequency",
                options=freq_options,
                index=freq_options.index(default_freq),
                key="gardening_frequency"
            )
            info['gardening_frequency'] = frequency

            # → New: Day of the Week
            days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday", "Not Specified"]
            default_day = info.get('gardening_day_of_week', days[0])
            day = st.selectbox(
                "Gardening Day of the Week",
                options=days,
                index=days.index(default_day),
                key="gardening_day_of_week"
            )
            info['gardening_day_of_week'] = day

            info['gardening_access'] = st.text_input(
                "Access Method for Gardeners",
                value=info.get('gardening_access', '')
            )
            info['gardening_finish_steps'] = st.text_area(
                "Post-Gardening Procedures",
                value=info.get('gardening_finish_steps', '')
            )
            info['gardening_identity_confirmation'] = st.text_area(
                "Gardening Crew Identity Verification",
                value=info.get('gardening_identity_confirmation', '')
            )

        # --- Pool Maintenance Service ---
        if "Pool Maintenance" in selected_services:
            st.subheader("🏊 Pool Maintenance Info")
            info = st.session_state.convenience_seeker_info

            info['pool_name'] = st.text_input(
                "Pool Maintenance Company Name",
                value=info.get('pool_name', '')
            )
            info['pool_number'] = st.text_input(
                "Pool Company Phone Number",
                value=info.get('pool_number', '')
            )

            # → Seasonal Months
            months = [
                "January","February","March","April","May","June",
                "July","August","September","October","November","December"
            ]
            default_months = info.get('pool_seasonal_months', [])

            selected_months = st.segmented_control(
                "Seasonal Months (select all that apply):",
                options=months,
                selection_mode="multi",
                default=default_months,
                key="pool_seasonal_months"
                )
            info['pool_seasonal_months'] = selected_months

            # → Seasonal Frequency
            freq_options = ["Monthly", "Bi-Weekly", "Weekly"]
            default_freq = info.get('pool_seasonal_frequency', freq_options[0])
            pool_freq = st.selectbox(
                "Seasonal Frequency:",
                options=freq_options,
                index=freq_options.index(default_freq),
                key="pool_seasonal_frequency"
            )
            info['pool_seasonal_frequency'] = pool_freq

            # → Day of the Week
            days = ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"]
            default_day = info.get('pool_day_of_week', days[0])
            pool_day = st.selectbox(
                "Pool Maintenance Day of the Week:",
                options=days,
                index=days.index(default_day),
                key="pool_day_of_week"
            )
            info['pool_day_of_week'] = pool_day

            info['pool_access'] = st.text_input(
                "Access Method for Pool Techs",
                value=info.get('pool_access', '')
            )
            info['pool_finish_steps'] = st.text_area(
                "Post-Maintenance Procedures",
                value=info.get('pool_finish_steps', '')
            )
            info['pool_identity_confirmation'] = st.text_area(
                "Pool Crew Identity Verification",
                value=info.get('pool_identity_confirmation', '')
            )

    # --- Save Button ---
    if st.button("💾 Quality-Oriented Household Services Info"):
        st.session_state["convenience_seeker_saved"] = True
        st.success("✅ Services information saved successfully!")


def rent_own():
    st.write("🏠 Home Ownership Status")

    if "rent_own_info" not in st.session_state:
        st.session_state.rent_own_info = {}

    housing_status = st.radio(
        "Do you rent or own your home?",
        options=["Rent", "Own"],
        index=0 if st.session_state.get("housing_status", "") == "Rent" else 1,
        key="housing_status"
    )   

    st.session_state.rent_own_info["housing_status"] = housing_status

    if housing_status == "Rent":
        st.subheader("🏢 Property Management Info")

        st.session_state.rent_own_info["property_management_name"] = st.text_input("Company Name")
        st.session_state.rent_own_info["property_management_number"] = st.text_input("Company Phone Number")
        st.session_state.rent_own_info["property_management_email"] = st.text_input("Company Email")
        st.session_state.rent_own_info["property_management_description"] = st.text_area(
            "When to Contact", placeholder="E.g. Roof issues, leaking pipe, parking, etc."
        )

    elif housing_status == "Own":
        st.subheader("🧰 Homeowner Contacts")

        homeowner_contacts_options = st.segmented_control(
            "Which service contacts are applicable?",
            options=["Handyman/Contractor", "Electrician", "Exterminator", "Plumber", "HOA"],
            selection_mode="multi",
            default=st.session_state.get("homeowner_contacts_options", []),
            key="homeowner_contacts_options"
        )

        st.session_state.rent_own_info["homeowner_contacts_options"] = homeowner_contacts_options

        # Utility function for section layout
        def contact_section(role):
            st.write(f"### {role}")
            name = st.text_input(f"{role} Name")
            number = st.text_input(f"{role} Phone Number")
            description = st.text_area(f"When to Contact {role}?")
            if name: st.session_state.rent_own_info[f"{role.lower()}_name"] = name
            if number: st.session_state.rent_own_info[f"{role.lower()}_number"] = number
            if description: st.session_state.rent_own_info[f"{role.lower()}_description"] = description

        if "Handyman/Contractor" in homeowner_contacts_options:
            contact_section("Handyman")

        if "Electrician" in homeowner_contacts_options:
            contact_section("Electrician")

        if "Exterminator" in homeowner_contacts_options:
            contact_section("Exterminator")

        if "Plumber" in homeowner_contacts_options:
            contact_section("Plumber")

        if "HOA" in homeowner_contacts_options:
            st.write("🏘️ HOA / Property Management")

            st.session_state.rent_own_info["property_management_name"] = st.text_input("Company Name (HOA)")
            st.session_state.rent_own_info["property_management_number"] = st.text_input("Phone Number (HOA)")
            st.session_state.rent_own_info["property_management_email"] = st.text_input("Email (HOA)")
            st.session_state.rent_own_info["property_management_description"] = st.text_area(
                "When to Contact (HOA)",
                placeholder="E.g. roof issues, bylaws, common areas, etc."
            )
        # --- Save Button ---
    if st.button("💾 Save Housing Status & Contacts Info"):
        st.session_state["rent_own_saved"] = True
        st.success("✅ Housing Status and contact information saved successfully!")


def security_convenience_ownership():
    st.subheader("Level 4: Home Security, Privacy, Quality-Orientation, and Support")
    # Step 1: User Input
    home_security()
    convenience_seeker()
    rent_own()
    
    # Step 2: Preview prompt

    # Move this outside the expander
    user_confirmation = st.checkbox("✅ Confirm AI Prompt")
    st.session_state["user_confirmation"] = user_confirmation # store confirmation in session

    st.session_state.progress["level_4_completed"] = True
    save_progress(st.session_state.progress)

    if user_confirmation:
        prompt_emergency = emergency_kit_utilities_runbook_prompt()
        prompt_mail_trash = mail_trash_runbook_prompt()
        prompt_home_caretaker = home_caretaker_runbook_prompt()
        st.session_state["prompt_emergency"] = prompt_emergency
        st.session_state["prompt_mail_trash"] = prompt_mail_trash
        st.session_state["prompt_home_caretaker"]= prompt_home_caretaker
    else:
        st.session_state["prompt_emergency"] = None
        st.session_state["prompt_mail_trash"] = None
        st.session_state["prompt_home_caretaker"]= None

# Show prompt in expander
    with st.expander("AI Prompt Preview (Optional)"):
        if st.session_state.get("prompt_emergency"):
            st.markdown("#### 🆘 Emergency + Utilities Prompt")
            st.code(st.session_state["prompt_emergency"], language="markdown")
        if st.session_state.get("prompt_mail_trash"):
            st.markdown("#### 📬 Mail + Trash Prompt")
            st.code(st.session_state["prompt_mail_trash"], language="markdown")
        if st.session_state.get("prompt_home_caretaker"):
            st.markdown("#### 💝 Home Protection + Services Prompt")
            st.code(st.session_state["prompt_home_caretaker"], language="markdown")

    # Step 3: Generate runbook using reusable function
   # st.write("Next, click the button to generate your personalized utilities emergency runbook document:")

    # st.markdown("### 🧪 Debug Info")

    # st.write("🔑 **API Key Loaded:**", "✅ Yes" if os.getenv("MISTRAL_TOKEN") else "❌ No")

    # st.write("✅ **User Confirmed Prompt:**", st.session_state.get("user_confirmation", False))

    # st.write("📄 **Emergency Prompt Exists:**", "✅ Yes" if st.session_state.get("prompt_emergency") else "❌ No")
    # st.code(st.session_state.get("prompt_emergency", "⚠️ Emergency prompt not generated."), language="markdown")

    # st.write("📬 **Mail & Trash Prompt Exists:**", "✅ Yes" if st.session_state.get("prompt_mail_trash") else "❌ No")
    # st.code(st.session_state.get("prompt_mail_trash", "⚠️ Mail/Trash prompt not generated."), language="markdown")

    generate_runbook_from_multiple_prompts(
        prompts=[
            st.session_state.get("prompt_emergency", ""),
            st.session_state.get("prompt_mail_trash", ""),
            st.session_state.get("prompt_home_caretaker", "")
        ],
        api_key=os.getenv("MISTRAL_TOKEN"),
        button_text="Complete Level 4 Mission",
        doc_heading="Comprehensive Housekeeping Runbook",
        doc_filename="housekeeping_runbook.docx"
    )
##### Level 5 - Emergency Kit Critical Documents

def emergency_kit_critical_documents():
    # Define categories and the corresponding documents
    documents = {
        'Identification Documents': [
            'Government-issued ID (Driver’s license, state ID)',
            'Social Security Card',
            'Birth Certificates',
            'Marriage/Divorce Certificates',
            'Citizenship/Immigration Documents',
            'Passport'
        ],
        'Health and Medical Documents': [
            'Health Insurance Cards',
            'Prescription Medications List',
            'Vaccination Records',
            'Emergency Medical Information',
            'Medical Power of Attorney'
        ],
        'Financial Documents': [
            'Bank Account Information',
            'Credit Cards/Debit Cards',
            'Checkbook',
            'Tax Returns (Last Year’s)',
            'Insurance Policies (Auto, Health, Home, Life, etc.)',
            'Investment Documents'
        ],
        'Homeownership or Rental Documents': [
            'Deed or Lease Agreement',
            'Mortgage or Rent Payment Records',
            'Home Insurance Policy'
        ],
         'Legal Documents': [
            'Will or Living Will',
            'Power of Attorney',
            'Property Title and Vehicle Titles',
            'Child Custody or Adoption Papers'
        ],
        'Emergency Contact Information': [
            'Contact List',
            'Emergency Plan'
        ],
        'Travel Documents': [
            'Travel Itinerary'
        ],
        'Educational Documents': [
            'School Records',
            'Diplomas and Degrees',
            'Certificates and Licenses'
        ],
        'Significant Documents': [
            'Pet Records',
            'Photos of Important Belongings',
            'Bankruptcy or Legal Filings'
        ]

    }

    # Initialize session state
    if "selected_documents" not in st.session_state:
        st.session_state.selected_documents = {}

    # 1) Category picker
    selected_category = st.selectbox(
        'Select a document category to view:',
        options=list(documents.keys()),
        key="selected_category"
    )

    # 2) Docs multiselect for that category
    if selected_category:
        st.write(f'You selected **{selected_category}**')

        # 2) Action buttons placed before the multiselect
        col1, col2 = st.columns(2)
        if col1.button('Add more categories', key="btn_add_more"):
            st.info("Pick another category above.")
        if col2.button('Finalize and Save All Selections', key="btn_finalize"):
            st.session_state.finalized = True

        # 3) Multi-select segmented control via horizontal checkboxes
        options = documents[selected_category]
        default = st.session_state.selected_documents.get(selected_category, [])
        cols = st.columns(len(options))
        new_picks = []
        for idx, opt in enumerate(options):
            # each checkbox lives in its own column
            checked = cols[idx].checkbox(
                opt,
                value=(opt in default),
                key=f"chk_{selected_category}_{idx}"
            )
            if checked:
                new_picks.append(opt)

        # save back
        st.session_state.selected_documents[selected_category] = new_picks

    # 5) If finalized, show all
    if st.session_state.get("finalized", False):
        st.header("✅ All Your Selections")
        for cat, docs in st.session_state.selected_documents.items():
            st.subheader(cat)
            for d in docs:
                st.write(f"• {d}")

def review_selected_documents():
    saved = st.session_state.get("selected_documents", {})
    if not saved:
        st.warning("No selections to review.")
        return

    st.header("📋 Review Selections")
    for cat, docs in saved.items():
        st.write(f"**{cat}:** {', '.join(docs)}")

    all_docs = [d for docs in saved.values() for d in docs]
    st.multiselect(
        "Tweak your list:",
        options=all_docs,
        default=all_docs,
        key="tweaked_docs"
    )
    if st.button("Save Tweaks", key="btn_save_tweaks"):
        st.success("Tweaks saved!")

def collect_document_details():
    selected = st.session_state.get("selected_documents", {})
    if not selected:
        st.warning("No documents selected. Go pick some first!")
        return

    # Initialize storage-confirmed flag
    if "storage_confirmed" not in st.session_state:
        st.session_state.storage_confirmed = False

    st.header("🗂 Document Access & Storage Details")

    PHYSICAL_STORAGE_OPTIONS = [
        "Canister","Closet","Drawer","Filing Cabinet","Handbag","Safe",
        "Safety Deposit Box","Storage Unit","Wallet","With Attorney", "With Financial Advisor/Accountant", "Other physical location"
    ]
    DIGITAL_STORAGE_OPTIONS = [
        "Computer/Tablet","Phone","USB flash drive","External hard drive",
        "Cloud storage (Google Drive, Dropbox, etc.)","Password Manager", "Mobile Application(s)", "Other digital location"
    ]

    # --- Step 0: Pick which storage types you use ---
    use_physical = st.checkbox("I use physical storage for my documents", key="use_physical")
    if use_physical:
        st.multiselect(
            "Select all physical storage locations you use:",
            options=PHYSICAL_STORAGE_OPTIONS,
            default=st.session_state.get("global_physical_storage", []),
            key="global_physical_storage"
        )

    use_digital = st.checkbox("I use digital storage for my documents", key="use_digital")
    if use_digital:
        st.multiselect(
            "Select all digital storage locations you use:",
            options=DIGITAL_STORAGE_OPTIONS,
            default=st.session_state.get("global_digital_storage", []),
            key="global_digital_storage"
        )

    # --- Step 0b: Confirm storage setups ---
    if st.button("Confirm storage types & locations", key="btn_confirm_storage"):
        errors = []
        if use_physical and not st.session_state.get("global_physical_storage"):
            errors.append("• select at least one physical storage location")
        if use_digital and not st.session_state.get("global_digital_storage"):
            errors.append("• select at least one digital storage location")
        if errors:
            st.error("Please:\n" + "\n".join(errors))
        else:
            st.session_state.storage_confirmed = True

    if not st.session_state.storage_confirmed:
        st.info("After selecting storage types & locations above, click **Confirm** to assign documents.")
        return

    # --- Step 1: Assign each document to chosen locations ---
    if "document_details" not in st.session_state:
        st.session_state.document_details = {}

    st.markdown("### Assign each document to one or more storage locations")
    all_assigned = True
    missing = []

    for category, docs in selected.items():
        if not docs:
            continue

        # Wrap the entire category in an expander
        with st.expander(category, expanded=False):
            for doc in docs:
                details = st.session_state.document_details.setdefault(doc, {})

                # Build the options from global storage lists
                options = []
                if use_physical:
                    options += st.session_state["global_physical_storage"]
                if use_digital:
                    options += st.session_state["global_digital_storage"]

                st.markdown(f"📄 **{doc}** — assign storage:")

                # Horizontal checkboxes, 4 per row
                picked = []
                for start in range(0, len(options), 4):
                    chunk = options[start : start + 4]
                    cols = st.columns(len(chunk))
                    for idx, opt in enumerate(chunk):
                        was = details.get("assigned_storage", [])
                        checked = cols[idx].checkbox(
                            opt,
                            value=(opt in was),
                            key=f"assign_{doc}_chk_{start+idx}"
                        )
                        if checked:
                            picked.append(opt)

                details["assigned_storage"] = picked

                if not picked:
                    all_assigned = False
                    missing.append(doc)

    # --- Step 2: Enforce that every document got assigned ---
    if not all_assigned:
        st.error("Please assign storage for all documents:")
        st.write(", ".join(missing))
        return

    # --- Step 3: Final save button ---
    if st.button("Save all document details", key="btn_save_details"):
        st.success("✅ All document details saved!")

    # Step 4: Ask storage-location questions
    st.header("🔍 Storage Location Details")

    # PHYSICAL STORAGE
    for storage in st.session_state.get("global_physical_storage", []):
        # normalize key name
        key_base = storage.lower().replace(" ", "_")
        with st.expander(f"{storage} Details", expanded=False):
            if storage == "Safety Deposit Box":
                st.text_input(
                    "Branch name & address:",
                    key=f"{key_base}_branch_address"
                )
                st.text_area(
                    "Emergency authorization required to retrieve contents:",
                    key=f"{key_base}_authorization"
                )

            elif storage == "Safe":
                st.text_input(
                    "Designated safe location (building/room/area):",
                    key=f"{key_base}_location"
                )
                st.text_area(
                    "Emergency steps & credentials needed to open safe:",
                    key=f"{key_base}_access_steps"
                )

            elif storage == "Storage Unit":
                st.text_input(
                    "Business name & address of unit:",
                    key=f"{key_base}_business_address"
                )
                st.text_area(
                    "Emergency authorization required for unit access:",
                    key=f"{key_base}_authorization"
                )

            elif storage == "With Attorney":
                st.text_area(
                    "Emergency contact method and proof of authorization needed:",
                    key=f"{key_base}_attorney_instructions"
                )

            elif storage == "Canister":
                st.text_input(
                    "Primary Canister location (building/room/cabinet/shelf):",
                    key=f"{key_base}_location"
                )
                st.text_area(
                    "Emergency steps & credentials needed to open canister:",
                    key=f"{key_base}_access_steps"
                )
                st.text_area(
                    "If secondary canisters are used , list each canister name & its location & contents:",
                    key=f"{key_base}_contents"
                )

            elif storage == "Drawer":
                st.text_input(
                    "Primary Drawer location (building/room/cabinet):",
                    key=f"{key_base}_location"
                )
                st.text_area(
                    "Emergency steps & credentials needed to open drawer:",
                    key=f"{key_base}_access_steps"
                )
                st.text_area(
                    "If secondary drawers are used , list each drawer name & its location & contents:",
                    key=f"{key_base}_contents"
                )

            elif storage == "Filing Cabinet":
                st.text_input(
                    "Primary Filing cabinet location (building/room/identifier):",
                    key=f"{key_base}_location"
                )
                st.text_area(
                    "Emergency steps & credentials needed to open cabinet:",
                    key=f"{key_base}_access_steps"
                )
                st.text_area(
                    "If secondary filing cabinets are used, list each cabinet name & its location & contents:",
                    key=f"{key_base}_contents"
                )

            elif storage == "Wallet":
                st.text_input(
                    "Wallet location (building/room/drawer/closet/bag):",
                    key=f"{key_base}_location"
                )
                st.text_area(
                    "Emergency steps & credentials to retrieve wallet:",
                    key=f"{key_base}_access_steps"
                )
                st.text_area(
                    "If secondary wallets are used, list each & its location & contents:",
                    key=f"{key_base}_contents"
                )

            elif storage == "Handbag":
                st.text_input(
                    "Primary Handbag location (building/room/drawer/closet):",
                    key=f"{key_base}_location"
                )
                st.text_area(
                    "Emergency steps & credentials to retrieve handbag:",
                    key=f"{key_base}_access_steps"
                )
                st.text_area(
                    "If secondary handbags are used, list its location & contents:",
                    key=f"{key_base}_contents"
                )

            elif storage == "Other physical location":
                st.text_input(
                    "Other location description (building/room/address):",
                    key=f"{key_base}_location"
                )
                st.text_area(
                    "Emergency steps & credentials to access this location:",
                    key=f"{key_base}_access_steps"
                )
                st.text_area(
                    "If multiple, list each location & its contents:",
                    key=f"{key_base}_contents"
                )

    # DIGITAL STORAGE
    for storage in st.session_state.get("global_digital_storage", []):
        key_base = storage.lower().replace(" ", "_").replace("/", "_")
        with st.expander(f"{storage} Details", expanded=False):
            if storage in ["Computer/Tablet", "Phone"]:
                st.text_input(
                    "Designated place (room/surface/storage)primary device:",
                    key=f"{key_base}_location"
                )
                st.text_area(
                    "Emergency steps & credentials to access device:",
                    key=f"{key_base}_access_steps"
                )
                st.text_area(
                    "If secondary devices exists, list each, its location and contents:",
                    key=f"{key_base}_contents"
                )

            elif storage in ["USB flash drive", "External hard drive"]:
                st.text_input(
                    "Designated place (room/surface/storage) for primary device:",
                    key=f"{key_base}_location"
                )
                st.text_area(
                    "Emergency steps & credentials to access drive:",
                    key=f"{key_base}_access_steps"
                )
                st.text_area(
                    "If secondary devices exists, list each & its location and contents:",
                    key=f"{key_base}_contents"
                )

            elif storage == "Cloud storage (Google Drive, Dropbox, etc.)":
                st.text_input(
                    "Primary Cloud platform name & link:",
                    key=f"{key_base}_platform"
                )
                st.text_area(
                    "Emergency steps & credentials to access account:",
                    key=f"{key_base}_access_steps"
                )
                st.text_area(
                    "If secondary platforms, list each & its link and contents:",
                    key=f"{key_base}_contents"
                )

            elif storage == "Password Manager":
                st.text_input(
                    "Password manager name:",
                    key=f"{key_base}_platform"
                )
                st.text_area(
                    "Emergency steps & credentials to access vault:",
                    key=f"{key_base}_access_steps"
                )
                st.text_area(
                    "If multiple vaults, list each & its contents:",
                    key=f"{key_base}_contents"
                )

            elif storage == "Mobile Application(s)":
                # New Mobile Application questions
                st.text_area(
                    "If multiple mobile applications are used, name each and note what is stored in each:",
                    key=f"{key_base}_apps_and_contents"
                )
                st.text_area(
                    "In an emergency, what steps and credentials are required for someone else to access the mobile application accounts holding key documents?:",
                    key=f"{key_base}_access_steps"
                )

    # Merge all storage‐location keys into document_details
    for doc, details in st.session_state.document_details.items():
        # for every storage the user selected
        for storage in st.session_state.get("global_physical_storage", []) \
                       + st.session_state.get("global_digital_storage", []):
            key_base = storage.lower().replace(" ", "_").replace("/", "_")
            # list every suffix you might have used
            for suffix in [
                "branch_address", "authorization",
                "location", "business_address",
                "attorney_instructions",
                "access_steps", "contents",
                "apps_and_contents", "platform"
            ]:
                full_key = f"{key_base}_{suffix}"
                if full_key in st.session_state:
                    # copy it into the per‐doc details dict
                    details[full_key] = st.session_state[full_key]

    # Final Save
    if st.button("Save all document & storage details", key="btn_save_all"):
        st.success("✅ All details saved!")


def generate_kit_tab():
    """Renders the Generate Kit UI and uses generate_runbook_from_prompt to run the LLM and export."""
    st.header("📦 Generate Emergency Document Kit")

    # 1) Build and show the prompt (optional—you can hide this if you don't want the user to see it)
    prompt = emergency_kit_document_prompt()
    with st.expander("Preview LLM prompt", expanded=False):
        st.code(prompt, language="markdown")

    # 2) Ask the user to confirm before sending
    st.checkbox("✅ I confirm this prompt is correct", key="user_confirmation")

    # 3) Delegate to your reusable runbook function
    generate_runbook_from_prompt(
        prompt=prompt,
        api_key=os.getenv("MISTRAL_TOKEN"),
        button_text="Generate Emergency Kit Runbook",
        doc_heading="Emergency Document Kit",
        doc_filename="emergency_document_kit.docx"
    )

##### Bonus - Additional Instructions for Guest/House Sitters

def bonus_level():
    st.write("🎁 Bonus Level")

    # ─── Initialize session_state keys ────────────────────────
    st.session_state.setdefault('bonus_info', {})
    st.session_state.setdefault('prompt_emergency', None)
    st.session_state.setdefault('prompt_bonus', None)
    st.session_state.setdefault('prompt_mail_trash', None)
    st.session_state.setdefault('prompt_home_caretaker', None)

    # Confirmation flag for generation
    st.session_state.setdefault('bonus_generate_confirm', False)

    # Ensure progress flags exist
    st.session_state.progress.setdefault("level_2_completed", False)
    st.session_state.progress.setdefault("level_3_completed", False)
    st.session_state.progress.setdefault("level_4_completed", False)

    # ─── Create two tabs ──────────────────────────────────────
    tab1, tab2 = st.tabs([
        "1️⃣ Bonus Input",
        "2️⃣ Generate Runbook"
    ])

    # ─── Tab 1: Collect Bonus Inputs ─────────────────────────
    with tab1:
        info = st.session_state.bonus_info

        with st.expander("🏠 Home Maintenance", expanded=True):
            info['maintenance_tasks'] = st.text_area(
                "List regular home maintenance tasks (e.g., changing bulbs, checking smoke detectors, cleaning filters):",
                value=info.get('maintenance_tasks', '')
            )
            info['appliance_instructions'] = st.text_area(
                "Instructions for operating/maintaining major appliances and systems:",
                value=info.get('appliance_instructions', '')
            )

        with st.expander("📋 Home Rules & Preferences", expanded=False):
            info['house_rules'] = st.text_area(
                "Guest/house­sitter rules or preferences:",
                value=info.get('house_rules', '')
            )
            info['cultural_practices'] = st.text_area(
                "Cultural/religious practices guests should be aware of:",
                value=info.get('cultural_practices', '')
            )

        with st.expander("🧹 Housekeeping & Cleaning", expanded=False):
            info['housekeeping_instructions'] = st.text_area(
                "Basic housekeeping/cleaning routines and supply locations:",
                value=info.get('housekeeping_instructions', '')
            )
            info['cleaning_preferences'] = st.text_area(
                "Specific cleaning preferences or routines:",
                value=info.get('cleaning_preferences', '')
            )

        with st.expander("🎮 Entertainment & Technology", expanded=False):
            info['entertainment_info'] = st.text_area(
                "How to operate entertainment systems and streaming services:",
                value=info.get('entertainment_info', '')
            )
            info['device_instructions'] = st.text_area(
                "Instructions for using/charging personal devices:",
                value=info.get('device_instructions', '')
            )

        if st.button("💾 Save Bonus Info"):
            st.success("✅ Bonus level information saved!")

    # ─── Tab 2: Generate Runbook ─────────────────────────────
        with tab2:
            st.subheader("Select and Generate Your Runbook")

            # 1) Must have at least one Bonus input
            bonus_info = st.session_state.bonus_info
            if not any(v and str(v).strip() for v in bonus_info.values()):
                st.error("🔒 Please complete at least one Bonus section in Tab 1 before proceeding.")
                return

            # 2) Mission choices
            missions = [
                "Bonus Level Mission",
                "Mail & Trash + Bonus Mission",
                "Full Runbook Mission"
            ]
            choice = st.radio("Which runbook would you like to generate?", options=missions, key="bonus_runbook_choice")

            # 3) Confirmation checkbox
            confirmed = st.checkbox(
                "✅ Confirm AI Prompt",
                value=st.session_state.get("user_confirmation", False),
                key="user_confirmation"
            )
            if not confirmed:
                st.info("Please confirm to preview and generate your runbook.")
                return

            # 4) Now that user is ready, enforce prerequisites
            if choice == missions[0] and not st.session_state.progress["level_2_completed"]:
                st.warning("🔒 Complete Level 2 before generating the Bonus Level runbook.")
                return
            if choice == missions[1] and not st.session_state.progress["level_3_completed"]:
                st.warning("🔒 Complete Level 3 before generating the Mail & Trash + Bonus runbook.")
                return
            if choice == missions[2] and not st.session_state.progress["level_4_completed"]:
                st.warning("🔒 Complete Level 4 before generating the Full runbook.")
                return

            # 5) Build all prompts
            st.session_state.prompt_emergency      = emergency_kit_utilities_runbook_prompt()
            st.session_state.prompt_bonus          = bonus_level_runbook_prompt()
            st.session_state.prompt_mail_trash     = mail_trash_runbook_prompt()
            st.session_state.prompt_home_caretaker = home_caretaker_runbook_prompt()

            # 6) Assemble the exact prompts, labels, and filenames
            if choice == missions[0]:
                prompts     = [st.session_state.prompt_emergency, st.session_state.prompt_bonus]
                labels      = ["🆘 Emergency + Utilities Prompt", "🎁 Bonus Level Prompt"]
                button_text = "Complete Bonus Level Mission"
                doc_heading = "Home Emergency Runbook with Bonus Level"
                doc_file    = "home_runbook_with_bonus.docx"
            elif choice == missions[1]:
                prompts     = [st.session_state.prompt_emergency, st.session_state.prompt_mail_trash, st.session_state.prompt_bonus]
                labels      = ["🆘 Emergency + Utilities Prompt", "📫 Mail & Trash Prompt", "🎁 Bonus Level Prompt"]
                button_text = "Complete Mail & Trash + Bonus Mission"
                doc_heading = "Emergency + Mail & Trash Runbook with Bonus"
                doc_file    = "runbook_mail_trash_bonus.docx"
            else:
                prompts     = [
                    st.session_state.prompt_emergency,
                    st.session_state.prompt_mail_trash,
                    st.session_state.prompt_home_caretaker,
                    st.session_state.prompt_bonus
                ]
                labels      = [
                    "🆘 Emergency + Utilities Prompt",
                    "📫 Mail & Trash Prompt",
                    "💝 Home Services Prompt",
                    "🎁 Bonus Level Prompt"
                ]
                button_text = "Complete Full Mission"
                doc_heading = "Complete Emergency Runbook with Bonus and Services"
                doc_file    = "runbook_full_mission.docx"

            # 7) Preview selected prompts
            for lbl, p in zip(labels, prompts):
                st.markdown(f"#### {lbl}")
                st.code(p, language="markdown")

            # 8) Generate runbook button
            generate_runbook_from_multiple_prompts(
                prompts=prompts,
                api_key=os.getenv("MISTRAL_TOKEN"),
                button_text=button_text,
                doc_heading=doc_heading,
                doc_filename=doc_file
            )

### Call App Functions
if __name__ == "__main__":
    main()